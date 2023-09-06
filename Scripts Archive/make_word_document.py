from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image

def find_all_headers(doc):

    # List to store all the header texts
    headers_list = []

    # Loop through all paragraphs in the document
    for paragraph in doc.paragraphs:
        # Check if the paragraph is a header (Heading style) and add it to the list
        if paragraph.style.name.startswith('Heading'):
            headers_list.append(paragraph.text)

    return headers_list

def print_content_between_headers(doc, start_header):
    '''Currently only seeming to print the paragaph following the header but nothing more (not all content inbetween headers)'''
    # Find all the headers in the document
    headers_list = find_all_headers(doc)

    # Find the index of the start header in the list
    try:
        start_index = headers_list.index(start_header)
    except ValueError:
        print("Start header not found.")
        return

    # Find the index of the next header after the start header
    end_index = start_index + 1

    # Print the content between the start header and the next header
    if end_index < len(headers_list):
        end_header = headers_list[end_index]
        print(f"Content between '{start_header}' and '{end_header}':")
        found_start_header = False

        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading') and paragraph.text == start_header:
                found_start_header = True
                continue

            if paragraph.style.name.startswith('Heading') and paragraph.text == end_header:
                break

            if found_start_header:
                print(paragraph.text)
                return paragraph.text
    else:
        print(f"Content after '{start_header}':")
        found_start_header = False

        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading') and paragraph.text == start_header:
                found_start_header = True
                continue

            if found_start_header:
                print(paragraph.text)
                return paragraph.text


def get_inline_shape_type(inline_shape):
    if inline_shape.type == 3:
        return "Picture"
    elif inline_shape.type == 7:
        return "Chart"
    else:
        return "Other"

def print_inline_shapes_breakdown(doc):

    inline_shapes_count = len(doc.inline_shapes)
    print(f"Total number of inline shapes: {inline_shapes_count}\n")

    inline_shapes_by_type = {"Picture": 0, "Chart": 0, "Other": 0}

    for inline_shape in doc.inline_shapes:
        shape_type = get_inline_shape_type(inline_shape)
        inline_shapes_by_type[shape_type] += 1

    for shape_type, count in inline_shapes_by_type.items():
        print(f"{shape_type}s: {count}")
    

def add_image_after_paragraph(doc, image_path, target_paragraph_text):
    # Find the target paragraph
    target_paragraph = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text == target_paragraph_text:
            target_paragraph = paragraph
            break

    if target_paragraph is None:
        raise ValueError(f"Paragraph with text '{target_paragraph_text}' not found in the document.")

    # Calculate the page width in twips (1 inch = 1440 twips)
    page_width_twips = Inches(8.5).emu

    # Calculate the left position to center the image
    image_width = Inches(5)  # Set the width of the image as desired
    left_twips = (page_width_twips - image_width.emu) // 2

    # Load the image to get its original dimensions
    image = Image.open(image_path)
    original_width, original_height = image.size

    # Calculate the aspect ratio of the image
    aspect_ratio = original_width / original_height

    # Calculate the desired image width and height
    desired_width = Inches(5)  # Set the desired width of the image
    desired_height = None  # Set to None to maintain aspect ratio

    # Calculate the new height based on the desired width and aspect ratio
    if desired_height is None:
        desired_height = desired_width / aspect_ratio

    # Calculate the new width based on the desired height and aspect ratio
    if desired_width is None:
        desired_width = desired_height * aspect_ratio

    # Create a new paragraph to insert the image
    new_paragraph = doc.add_paragraph("")

    # Create a new run to insert the image
    run = new_paragraph.add_run()

    # Add the image to the run with the adjusted scale
    run.add_picture(image_path, width=desired_width, height=desired_height)

    # Center align the image within the paragraph
    new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Usage example:
input_doc_path = "80 Queen St - Analytics Report - Template.docx"
output_doc_path = "80 Queen St - Analytics Report - Template2.docx"

image_to_add_path = '80Q - Water Meters/From EBI/Plot Data/North Face(M9).png'
start_header = "Water Usage"  # Replace this with the actual start header
doc = Document(input_doc_path)
paragraph_text = print_content_between_headers(doc, start_header)
print_inline_shapes_breakdown(doc)

add_image_after_paragraph(doc, image_to_add_path, paragraph_text)

doc.save(output_doc_path)